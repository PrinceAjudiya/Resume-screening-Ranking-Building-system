from preprocessing import docx_processing  as doc, textract_processing as txt
from text_processing import tf_idf_cosine_similarity as tf_idf, doc2vec_comparison as d2v
from text_processing import cv_cosine_similarity as cv
import os
import pathlib
import pdfplumber
import PyPDF2
import docx
import PyPDF2
def process_files(req_document, resume_docs):

    resume_doc_text = []

    for doct in resume_docs:
        print("File ----------------------------------", doct)
        # file_extension = pathlib.Path(resume_docs[doct]).suffix
        split_tup = os.path.splitext(doct)
        file_extension = split_tup[1]

        print("Helo.........................", file_extension)
        if file_extension == '.pdf':

            filename = os.path.basename(doct)

            split_tup = os.path.splitext(filename)

            # extract the file name and extension
            file_name = split_tup[0]

            print("File name-----------------------------", file_name)
            print("Hello if running-----------------------------------------")

            mydoc = docx.Document()
            pdfFileObj = open(doct, 'rb')  # pdffile loction
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

            with pdfplumber.open(pdfFileObj) as pdf:
                for i in range(0, pdfReader.numPages):
                    first_page = pdf.pages[i]
                    pdfContent = first_page.extract_text()
                    mydoc.add_paragraph(pdfContent)

            fp = "D:\\sem-7\\SGP\\Resume Ranking\\static\\UPLOAD_FOLDER\\" + file_name + ".docx"

            mydoc.save(
                "D:\\sem-7\\SGP\\Resume Ranking\\static\\UPLOAD_FOLDER\\" + file_name + ".docx")

            resume_doc_text.append(txt.get_content_as_string(fp))

        else:
            print("Hello else running---------------------------------------------")
            resume_doc_text.append(txt.get_content_as_string(doct))



    split_tup = os.path.splitext(req_document)
    file_extension_req = split_tup[1]

    if file_extension_req == ".pdf":
        filename_req = os.path.basename(req_document)

        split_tup = os.path.splitext(filename_req)

        # extract the file name and extension
        file_name_req = split_tup[0]

        print("File name-----------------------------", file_name_req)

        mydoc2 = docx.Document()
        pdfFileObj = open(req_document, 'rb')  # pdffile loction
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        with pdfplumber.open(pdfFileObj) as pdf:
            for i in range(0, pdfReader.numPages):
                first_page = pdf.pages[i]
                pdfContent = first_page.extract_text()
                mydoc2.add_paragraph(pdfContent)

        fp = "D:\\sem-7\\SGP\\Resume Ranking\\static\\UPLOAD_FOLDER\\" + file_name_req + ".docx"

        mydoc2.save(
            "D:\\sem-7\\SGP\\Resume Ranking\\static\\UPLOAD_FOLDER\\" + file_name_req + ".docx")

        req_doc_text = txt.get_content_as_string(fp)


    else:
        req_doc_text = txt.get_content_as_string(req_document)

    cos_sim_list = tf_idf.get_tf_idf_cosine_similarity(req_doc_text, resume_doc_text)
    final_doc_rating_list = []
    zipped_docs = zip(cos_sim_list, resume_docs)
    sorted_doc_list = sorted(zipped_docs, key=lambda x: x[0], reverse=True)
    for element in sorted_doc_list:
        doc_rating_list = []
        doc_rating_list.append(os.path.basename(element[1]))
        doc_rating_list.append("{:.0%}".format(element[0]))
        final_doc_rating_list.append(doc_rating_list)
    return final_doc_rating_list
