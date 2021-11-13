# #from pdf2docx import extract_tables
# from pdf2docx import parse
#
#
# path = 'D:\\sem-7\\SGP\\Resume Ranking\\AI_ML_Learning-master\\AI_ML_Learning-master\\resume_rating\\static\\UPLOAD_FOLDER\\Prince Ajudiya resume.pdf'
# path2 = 'D:'
#
#
# parse(path, path2, start=1, end=2)
#
# import os
# import subprocess
#
# for top, dirs, files in os.walk('D:\\sem-7\\SGP\\Resume Ranking\\AI_ML_Learning-master\\AI_ML_Learning-master\\resume_rating\\static\\UPLOAD_FOLDER'):
#     for filename in files:
#         if filename.endswith('.pdf'):
#             print("Hello")
#             abspath = os.path.join(top, filename)
#             subprocess.call('lowriter --invisible --convert-to doc "{}"'
#                             .format(abspath), shell=True)

import PyPDF2
import os
import docx
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO

mydoc = docx.Document()  # document type
pdfFileObj = open('D:\\sem-7\\SGP\\Resume Ranking\\AI_ML_Learning-master\\AI_ML_Learning-master\\resume_rating\\static\\UPLOAD_FOLDER\\18IT019_vandita_chapadia.pdf', 'rb')  # pdffile loction
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)  # define pdf reader object
rsrcmgr = PDFResourceManager()
retstr = StringIO()
codec = 'utf-8'
laparams = LAParams()
device = TextConverter(rsrcmgr, retstr, laparams=laparams)
fp = open('D:\\sem-7\\SGP\\Resume Ranking\\AI_ML_Learning-master\\AI_ML_Learning-master\\resume_rating\\static\\UPLOAD_FOLDER\\18IT019_vandita_chapadia.pdf', 'rb')
interpreter = PDFPageInterpreter(rsrcmgr, device)

# Loop through all the pages

for pageNum in range(0, pdfReader.numPages):
    print("Hello")
    pageObj = pdfReader.getPage(pageNum)
    pdfContent = pageObj.extractText()  # extracts the content from the page.
    print(pdfContent)  # print statement to test output in the terminal. codeline optional.
    mydoc.add_paragraph(pdfContent)  # this adds the content to the word document


# text = ''
# for page in PDFPage.get_pages(fp , check_extractable=True):
#         interpreter.process_page(page)
#
# text = retstr.getvalue()
# print(text)
# mydoc.add_paragraph(text)


mydoc.save("D:\\sem-7\\SGP\\Resume Ranking\\AI_ML_Learning-master\\AI_ML_Learning-master\\resume_rating\\static\\filename.docx")


